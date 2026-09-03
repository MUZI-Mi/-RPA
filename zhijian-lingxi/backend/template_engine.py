"""模板渲染与多格式导出

- fill_excel / fill_word：把 Excel/Word 模板里的 {{占位符}} 替换为数据，
  {{__rows__}} 标记处展开表格数据行
- to_xlsx / to_docx / to_pdf：无模板时直接由数据表生成文件
- PDF 使用 reportlab，需注册系统中文字体避免乱码

第三方库（openpyxl / python-docx / reportlab）延迟到方法内导入，
避免模块加载阶段因缺依赖而崩溃。
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

from config import TEMPLATE_DIR


class TemplateEngine:
    """模板填充与文档生成工具类。"""

    # ---- 模板查找 ----
    @staticmethod
    def find_template(ref: str) -> Optional[Path]:
        """按模板 name 或 id 查 templates 表，返回 TEMPLATE_DIR 内真实路径。"""
        from database import get_conn

        if not ref:
            return None
        with get_conn() as conn:
            row = conn.execute(
                "SELECT filename FROM templates WHERE name = ? OR id = ?", (ref, ref)
            ).fetchone()
        if not row:
            return None
        path = TEMPLATE_DIR / row["filename"]
        return path if path.exists() else None

    # ---- PDF 中文字体 ----
    @staticmethod
    def _register_cjk_font() -> str:
        """注册中文字体，返回可用字体名；找不到返回 Helvetica。"""
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            return "Helvetica"
        candidates = [
            (Path("C:/Windows/Fonts/simhei.ttf"), None),
            (Path("C:/Windows/Fonts/simsun.ttc"), 0),
            (Path("C:/Windows/Fonts/msyh.ttc"), 0),
            (Path("C:/Windows/Fonts/msyh.ttf"), None),
        ]
        for path, sub in candidates:
            if not path.exists():
                continue
            try:
                if sub is not None:
                    pdfmetrics.registerFont(TTFont("CJKFont", str(path), subfontIndex=sub))
                else:
                    pdfmetrics.registerFont(TTFont("CJKFont", str(path)))
                return "CJKFont"
            except Exception:
                continue
        return "Helvetica"

    # ---- Excel ----
    @staticmethod
    def _cell_value(v: Any) -> Any:
        """单元格值规范化：列表/字典等非标量转成易读字符串，避免 openpyxl 写入失败。"""
        if isinstance(v, (list, tuple)):
            return "、".join(str(x) for x in v) if v else "无"
        if isinstance(v, dict):
            return str(v)
        return v

    @staticmethod
    def fill_excel(
        template_path: Path, data: Dict[str, Any], rows: List[Dict], out_path: Path
    ) -> Path:
        """按模板填充：标量占位符替换 + {{__rows__}} 行展开。

        约定：{{__rows__}} 所在行写入表头，其下插入数据行（原模板该区域内容会被覆盖，
        模板作者需预留行数）。
        """
        from openpyxl import load_workbook

        wb = load_workbook(str(template_path))
        ws = wb.active
        row_anchor = None
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str) or "{{" not in cell.value:
                    continue
                if "{{__rows__}}" in cell.value:
                    row_anchor = cell.row
                    continue
                text = cell.value
                for k, v in (data or {}).items():
                    text = text.replace("{{%s}}" % k, "" if v is None else str(v))
                cell.value = text
        if row_anchor is not None and rows:
            headers = list(rows[0].keys())
            for j, h in enumerate(headers, start=1):
                ws.cell(row=row_anchor, column=j, value=h)
            for i, r in enumerate(rows, start=1):
                for j, h in enumerate(headers, start=1):
                    ws.cell(row=row_anchor + i, column=j, value=TemplateEngine._cell_value(r.get(h, "")))
        wb.save(str(out_path))
        return out_path

    @staticmethod
    def to_xlsx(rows: List[Dict], out_path: Path) -> Path:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "报表"
        if rows:
            headers = list(rows[0].keys())
            ws.append(headers)
            for r in rows:
                ws.append([TemplateEngine._cell_value(r.get(h, "")) for h in headers])
        wb.save(str(out_path))
        return out_path

    # ---- Word ----
    @staticmethod
    def _replace_in_paragraph(p: Any, data: Dict[str, Any]) -> None:
        """替换单个段落：python-docx 段落文字常被拆在多个 run 里，
        需合并文本→替换→写回首 run、清空其余，避免漏替换。"""
        full = "".join(run.text for run in p.runs)
        if "{{" not in full:
            return
        for k, v in (data or {}).items():
            full = full.replace("{{%s}}" % k, "" if v is None else str(v))
        if p.runs:
            p.runs[0].text = full
            for run in p.runs[1:]:
                run.text = ""

    @staticmethod
    def fill_word(template_path: Path, data: Dict[str, Any], out_path: Path) -> Path:
        from docx import Document

        doc = Document(str(template_path))
        for p in doc.paragraphs:
            TemplateEngine._replace_in_paragraph(p, data)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        TemplateEngine._replace_in_paragraph(p, data)
        doc.save(str(out_path))
        return out_path

    @staticmethod
    def to_docx(rows: List[Dict], out_path: Path) -> Path:
        from docx import Document

        doc = Document()
        doc.add_heading("报表", level=1)
        if rows:
            headers = list(rows[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = str(h)
            for r in rows:
                cells = table.add_row().cells
                for j, h in enumerate(headers):
                    cells[j].text = "" if r.get(h) is None else str(r.get(h, ""))
        doc.save(str(out_path))
        return out_path

    # ---- PDF ----
    @staticmethod
    def to_pdf(rows: List[Dict], out_path: Path) -> Path:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Table, TableStyle

        font_name = TemplateEngine._register_cjk_font()
        doc = SimpleDocTemplate(str(out_path), pagesize=A4)
        elements = []
        if rows:
            headers = list(rows[0].keys())
            data = [[str(h) for h in headers]]
            for r in rows:
                data.append(["" if r.get(h) is None else str(r.get(h, "")) for h in headers])
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            elements.append(table)
        else:
            elements.append(Paragraph("（空报表，无数据）", getSampleStyleSheet()["BodyText"]))
        doc.build(elements)
        return out_path

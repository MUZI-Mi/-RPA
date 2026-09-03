"""模板引擎单元测试（_test_template_engine.py）

覆盖最近修复的回归点：
- _cell_value 把列表/字典规范成可写单元格值（回归：compliance_issues 空列表
  导致 openpyxl 报 "Cannot convert [] to Excel"）
- to_xlsx / fill_excel 写入含列表列的数据不再抛错，且空列表落为「无」
- fill_word / to_docx / to_pdf 冒烟验证

运行：在 backend 目录执行 `python _test_template_engine.py`
"""
import sys
import tempfile
from pathlib import Path

from template_engine import TemplateEngine

PASS = 0
FAIL = 0
SKIP = 0


def check(name, cond, detail=""):
    global PASS, FAIL
    if cond:
        PASS += 1
        print("  [PASS] %s" % name)
    else:
        FAIL += 1
        print("  [FAIL] %s  %s" % (name, detail))


def skip(name):
    global SKIP
    SKIP += 1
    print("  [SKIP] %s" % name)


def test_cell_value():
    print("\n== _cell_value ==")
    check("空列表 → '无'", TemplateEngine._cell_value([]) == "无", repr(TemplateEngine._cell_value([])))
    check("非空列表 → 顿号连接", TemplateEngine._cell_value(["a", "b"]) == "a、b", repr(TemplateEngine._cell_value(["a", "b"])))
    check("元组 → 顿号连接", TemplateEngine._cell_value(("x",)) == "x", repr(TemplateEngine._cell_value(("x",))))
    check("字典 → str", TemplateEngine._cell_value({"k": 1}) == "{'k': 1}", repr(TemplateEngine._cell_value({"k": 1})))
    check("字符串原样", TemplateEngine._cell_value("abc") == "abc", repr(TemplateEngine._cell_value("abc")))
    check("数字原样", TemplateEngine._cell_value(5) == 5, repr(TemplateEngine._cell_value(5)))
    check("None 原样", TemplateEngine._cell_value(None) is None, repr(TemplateEngine._cell_value(None)))


def test_to_xlsx_with_list_values():
    print("\n== to_xlsx：含列表列（回归：Cannot convert [] to Excel）==")
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "out.xlsx"
        rows = [
            {"姓名": "张三", "compliance_issues": ["金额异常"], "金额": 5000},
            {"姓名": "李四", "compliance_issues": [], "金额": 3200},
        ]
        TemplateEngine.to_xlsx(rows, out)
        check("文件生成", out.exists(), str(out))
        from openpyxl import load_workbook

        ws = load_workbook(out).active
        check("表头正确", [c.value for c in ws[1]] == ["姓名", "compliance_issues", "金额"],
              str([c.value for c in ws[1]]))
        r1 = [c.value for c in ws[2]]
        check("行1 非空列表 → 顿号文本", r1 == ["张三", "金额异常", 5000], str(r1))
        r2 = [c.value for c in ws[3]]
        check("行2 空列表 → '无'", r2 == ["李四", "无", 3200], str(r2))


def test_to_xlsx_empty_rows():
    print("\n== to_xlsx：空数据 ==")
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "empty.xlsx"
        TemplateEngine.to_xlsx([], out)
        check("空表也生成有效文件", out.exists(), str(out))


def test_fill_excel():
    print("\n== fill_excel：占位符 + __rows__ 展开 ==")
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        from openpyxl import Workbook, load_workbook

        tpl = td / "tpl.xlsx"
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "{{街道}}低保名单汇总"
        ws["A3"] = "{{__rows__}}"
        wb.save(str(tpl))

        out = td / "out.xlsx"
        TemplateEngine.fill_excel(tpl, {"街道": "浦东新区"}, [
            {"姓名": "张三", "compliance_issues": []},
            {"姓名": "李四", "compliance_issues": ["金额异常"]},
        ], out)

        ws2 = load_workbook(out).active
        check("标量占位符替换", ws2["A1"].value == "浦东新区低保名单汇总", repr(ws2["A1"].value))
        check("表头写入锚点行", ws2["A3"].value == "姓名", repr(ws2["A3"].value))
        check("行1 空列表 → '无'", ws2["A4"].value == "张三" and ws2["B4"].value == "无",
              str([ws2["A4"].value, ws2["B4"].value]))
        check("行2 非空列表 → 顿号", ws2["A5"].value == "李四" and ws2["B5"].value == "金额异常",
              str([ws2["A5"].value, ws2["B5"].value]))


def test_fill_word():
    print("\n== fill_word：占位符替换 ==")
    try:
        import docx
    except ImportError:
        skip("python-docx 未安装")
        return
    with tempfile.TemporaryDirectory() as td:
        tpl = Path(td) / "tpl.docx"
        d = docx.Document()
        d.add_paragraph("{{街道}}低保名单")
        d.save(str(tpl))
        out = Path(td) / "out.docx"
        TemplateEngine.fill_word(tpl, {"街道": "浦东新区"}, out)
        d2 = docx.Document(str(out))
        check("Word 占位符替换", "浦东新区低保名单" in d2.paragraphs[0].text, repr(d2.paragraphs[0].text))


def test_to_docx():
    print("\n== to_docx：含列表列 ==")
    try:
        import docx
    except ImportError:
        skip("python-docx 未安装")
        return
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "out.docx"
        TemplateEngine.to_docx([{"姓名": "张三", "compliance_issues": []}], out)
        check("docx 生成", out.exists(), str(out))


def test_to_pdf():
    print("\n== to_pdf：有数据 + 空表 ==")
    try:
        import reportlab
    except ImportError:
        skip("reportlab 未安装")
        return
    with tempfile.TemporaryDirectory() as td:
        out1 = Path(td) / "a.pdf"
        TemplateEngine.to_pdf([{"姓名": "张三", "compliance_issues": []}], out1)
        check("pdf(有数据)生成", out1.exists(), str(out1))
        out2 = Path(td) / "b.pdf"
        TemplateEngine.to_pdf([], out2)
        check("pdf(空表)生成", out2.exists(), str(out2))


def main():
    test_cell_value()
    test_to_xlsx_with_list_values()
    test_to_xlsx_empty_rows()
    test_fill_excel()
    test_fill_word()
    test_to_docx()
    test_to_pdf()

    print("\n===== 结果: %d 通过, %d 失败, %d 跳过 =====" % (PASS, FAIL, SKIP))
    sys.exit(1 if FAIL else 0)


if __name__ == "__main__":
    main()
